# -*- coding: utf-8 -*-
"""
services/ocr/table_path.py · 2026-05-21 multi-format refactor

Direct table reader for Excel (.xlsx / .xls), CSV, and Word (.docx).
These formats have structure already — sending them through OCR would
destroy the column layout that makes GL files parseable in the first
place.

Contract:
    extract_from_table_file(file_bytes, filename) -> Layer1Result
        - Reads file structurally, builds one Page per sheet (Excel) or one
          Page per logical table (Word) or one Page (CSV).
        - Page.full_text = textual rendering of the grid (for fallback / debug)
        - Page.table_headers = list of column header strings (best-effort)
        - Page.table_rows = list of {header: cell} dicts
        - avg_confidence = 1.0 (direct read — no OCR uncertainty)

The downstream Layer 2 picks up Page.table_rows + table_headers and feeds
the grid into the per-document-type prompt (see layer2_structure._page_to_text).
"""

from __future__ import annotations

import csv
import io
import logging
import time
from typing import List

from .schemas import Layer1Result, Page

logger = logging.getLogger(__name__)


SUPPORTED_TABLE_EXTENSIONS = {
    ".xlsx", ".xls", ".xlsm",
    ".csv", ".tsv",
    ".docx", ".doc",
    ".txt",
}


def extract_from_table_file(file_bytes: bytes, filename: str) -> Layer1Result:
    """Dispatch to the right reader based on filename extension."""
    if not file_bytes:
        raise ValueError("table_path: empty file bytes")
    name = (filename or "").lower()

    t0 = time.time()
    if name.endswith((".xlsx", ".xlsm", ".xls")):
        pages = _read_excel(file_bytes)
        engine = "table_path_excel"
    elif name.endswith(".csv"):
        pages = _read_csv(file_bytes, delimiter=",")
        engine = "table_path_csv"
    elif name.endswith(".tsv"):
        pages = _read_csv(file_bytes, delimiter="\t")
        engine = "table_path_tsv"
    elif name.endswith(".docx"):
        pages = _read_docx(file_bytes)
        engine = "table_path_docx"
    elif name.endswith(".doc"):
        # .doc legacy binary — fall back to extracting visible text only
        pages = _read_doc_fallback(file_bytes)
        engine = "table_path_doc"
    elif name.endswith(".txt"):
        pages = _read_txt(file_bytes)
        engine = "table_path_txt"
    else:
        raise ValueError(
            f"table_path: unsupported extension for {filename!r}; "
            f"supported: {sorted(SUPPORTED_TABLE_EXTENSIONS)}"
        )

    elapsed_ms = int((time.time() - t0) * 1000)
    return Layer1Result(
        pages=pages,
        page_count=len(pages),
        elapsed_ms=elapsed_ms,
        engine=engine,
        dpi=0,
    )


# ============================================================
# Excel readers (openpyxl for xlsx, xlrd-style fallback for .xls)
# ============================================================
def _read_excel(file_bytes: bytes) -> List[Page]:
    try:
        import openpyxl
    except ImportError as e:
        raise ImportError(
            "table_path: openpyxl required for Excel files. "
            "Install: pip install openpyxl"
        ) from e

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    pages: List[Page] = []
    for sheet_idx, sheet_name in enumerate(wb.sheetnames, start=1):
        ws = wb[sheet_name]
        rows_iter = ws.iter_rows(values_only=True)
        # Find first non-empty row → headers
        headers: List[str] = []
        body_rows: List[List[str]] = []
        seen_header = False
        for row in rows_iter:
            cells = ["" if c is None else str(c) for c in row]
            if not any(c.strip() for c in cells):
                continue
            if not seen_header:
                headers = [_normalize_header(c) for c in cells]
                seen_header = True
            else:
                body_rows.append(cells)

        # Build dict rows (header→cell). When column count mismatches, pad/truncate.
        table_rows = []
        for r in body_rows:
            padded = (r + [""] * len(headers))[: len(headers)]
            table_rows.append({h: v for h, v in zip(headers, padded)})

        full_text = _render_grid_text(sheet_name, headers, body_rows)
        pages.append(
            Page(
                page_number=sheet_idx,
                width=0,
                height=0,
                full_text=full_text,
                avg_confidence=1.0,
                blocks=[],
                table_headers=headers,
                table_rows=table_rows,
            )
        )
    wb.close()

    if not pages:
        # Empty workbook
        pages = [Page(
            page_number=1, width=0, height=0,
            full_text="(empty workbook)",
            avg_confidence=1.0, blocks=[],
            table_headers=[], table_rows=[],
        )]
    return pages


# ============================================================
# CSV / TSV reader
# ============================================================
def _read_csv(file_bytes: bytes, delimiter: str) -> List[Page]:
    # Detect encoding — most Thai CSVs are UTF-8 (with BOM) or CP874
    text = _decode_bytes(file_bytes)
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = [r for r in reader]
    if not rows:
        return [Page(
            page_number=1, width=0, height=0,
            full_text="(empty csv)", avg_confidence=1.0,
            blocks=[], table_headers=[], table_rows=[],
        )]
    headers = [_normalize_header(c) for c in rows[0]]
    body = rows[1:]
    table_rows = []
    for r in body:
        if not any((c or "").strip() for c in r):
            continue
        padded = (r + [""] * len(headers))[: len(headers)]
        table_rows.append({h: v for h, v in zip(headers, padded)})
    full_text = _render_grid_text("csv", headers, body)
    return [Page(
        page_number=1, width=0, height=0,
        full_text=full_text, avg_confidence=1.0,
        blocks=[], table_headers=headers, table_rows=table_rows,
    )]


# ============================================================
# Word (.docx) reader — extracts both paragraphs and tables
# ============================================================
def _read_docx(file_bytes: bytes) -> List[Page]:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise ImportError(
            "table_path: python-docx required for .docx files. "
            "Install: pip install python-docx"
        ) from e

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
    # Build one Page that combines paragraphs + tables
    text_parts: List[str] = []
    text_parts.append("\n".join(paragraphs_text))

    # If document has tables, treat the FIRST one as the primary grid
    table_headers: List[str] = []
    table_rows: List[dict] = []
    if doc.tables:
        first = doc.tables[0]
        if first.rows:
            header_cells = [_normalize_header(c.text) for c in first.rows[0].cells]
            table_headers = header_cells
            for row in first.rows[1:]:
                cells = [c.text for c in row.cells]
                padded = (cells + [""] * len(table_headers))[: len(table_headers)]
                row_dict = {h: v for h, v in zip(table_headers, padded)}
                if any((v or "").strip() for v in row_dict.values()):
                    table_rows.append(row_dict)
        # Append all tables to full_text for completeness
        for ti, tbl in enumerate(doc.tables):
            text_parts.append(f"\n[table {ti+1}]")
            for row in tbl.rows:
                text_parts.append(" | ".join(c.text for c in row.cells))

    full_text = "\n".join(text_parts).strip() or "(empty docx)"
    return [Page(
        page_number=1, width=0, height=0,
        full_text=full_text, avg_confidence=1.0,
        blocks=[],
        table_headers=table_headers or None,
        table_rows=table_rows or None,
    )]


def _read_doc_fallback(file_bytes: bytes) -> List[Page]:
    # Legacy .doc not directly supported without antiword. Decode as best-effort
    # text; downstream Layer 2 will get whatever is salvageable.
    text = _decode_bytes(file_bytes)
    # Strip non-printable bytes
    text = "".join(ch for ch in text if ch.isprintable() or ch in "\n\t ")
    return [Page(
        page_number=1, width=0, height=0,
        full_text=text or "(legacy .doc — limited extraction)",
        avg_confidence=0.7,  # legacy format, not fully trustable
        blocks=[],
        table_headers=None,
        table_rows=None,
    )]


def _read_txt(file_bytes: bytes) -> List[Page]:
    text = _decode_bytes(file_bytes)
    return [Page(
        page_number=1, width=0, height=0,
        full_text=text,
        avg_confidence=1.0,
        blocks=[],
        table_headers=None,
        table_rows=None,
    )]


# ============================================================
# Helpers
# ============================================================
def _decode_bytes(b: bytes) -> str:
    """Decode CSV / TXT bytes — try common encodings in order."""
    for enc in ("utf-8-sig", "utf-8", "cp874", "tis-620", "cp1252", "latin-1"):
        try:
            return b.decode(enc)
        except UnicodeDecodeError:
            continue
    return b.decode("utf-8", errors="replace")


def _normalize_header(s) -> str:
    """Light normalization for column headers — strip whitespace, drop None."""
    if s is None:
        return ""
    s = str(s).strip()
    return s


def _render_grid_text(label: str, headers: List[str], rows: List[List[str]]) -> str:
    """Build a pipe-delimited textual rendering for fallback / debug display."""
    out = [f"[{label}]"]
    if headers:
        out.append(" | ".join(headers))
    for r in rows:
        out.append(" | ".join("" if c is None else str(c) for c in r))
    return "\n".join(out)
