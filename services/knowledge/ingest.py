"""Document ingest: extract text, normalize, split into chunks.

This is the deterministic core of P1, deliberately free of any database or
network dependency so it is unit-testable on its own. Chunking is character-
based because Thai is written without word spaces, so token boundaries are not
available without a tokenizer; we pack paragraphs up to a character budget and
fall back to an overlapping character window for paragraphs that exceed it.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from services.knowledge.models import Chunk, ParsedDocument

# File types we can read as plain text directly. cp874 covers Thai on Windows.
TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".csv", ".tsv", ".log", ".json"}
TEXT_DECODE_ORDER = ("utf-8-sig", "utf-8", "cp874", "latin-1")

DEFAULT_MAX_CHARS = 800
DEFAULT_OVERLAP = 100


class UnsupportedDocument(Exception):
    """Raised when a file type cannot (yet) be parsed into text."""


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return _decode(data)
    extractor = _BINARY_EXTRACTORS.get(suffix)
    if extractor is not None:
        return extractor(data)
    raise UnsupportedDocument(f"unsupported file type: {suffix or '(none)'}")


def _decode(data: bytes) -> str:
    for encoding in TEXT_DECODE_ORDER:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_docx(data: bytes) -> str:
    from docx import Document  # python-docx; lazy so plain-text ingest stays light

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        parts: list[str] = []
        for sheet in workbook.worksheets:
            parts.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    finally:
        workbook.close()


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if not text.strip():
        # Image-only (scanned) PDF: no extractable text layer -> belongs to OCR.
        raise UnsupportedDocument("PDF 无可提取文本(可能是扫描件,需走 OCR)")
    return text


_BINARY_EXTRACTORS = {
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".pdf": _extract_pdf,
}


def normalize_text(raw: str) -> str:
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    text = re.sub(r"\n{3,}", "\n\n", text)  # collapse runs of blank lines to one
    return text.strip()


def chunk_text(
    text: str, *, max_chars: int = DEFAULT_MAX_CHARS, overlap: int = DEFAULT_OVERLAP
) -> list[Chunk]:
    if max_chars <= 0:
        raise ValueError("max_chars must be positive")
    pieces: list[str] = []
    buffer = ""
    for paragraph in (p.strip() for p in re.split(r"\n{2,}", text) if p.strip()):
        if len(paragraph) > max_chars:
            if buffer:
                pieces.append(buffer)
                buffer = ""
            pieces.extend(_window_split(paragraph, max_chars, overlap))
        elif buffer and len(buffer) + 1 + len(paragraph) > max_chars:
            pieces.append(buffer)
            buffer = paragraph
        else:
            buffer = f"{buffer}\n{paragraph}" if buffer else paragraph
    if buffer:
        pieces.append(buffer)
    return [Chunk(ordinal=i, text=p, char_count=len(p)) for i, p in enumerate(pieces)]


def _window_split(text: str, max_chars: int, overlap: int) -> list[str]:
    step = max_chars - overlap if overlap < max_chars else max_chars
    return [text[i : i + max_chars] for i in range(0, len(text), step)]


def ingest_document(
    filename: str,
    data: bytes,
    *,
    max_chars: int = DEFAULT_MAX_CHARS,
    overlap: int = DEFAULT_OVERLAP,
) -> ParsedDocument:
    text = normalize_text(extract_text(filename, data))
    return ParsedDocument(
        filename=filename,
        text=text,
        chunks=chunk_text(text, max_chars=max_chars, overlap=overlap),
    )
